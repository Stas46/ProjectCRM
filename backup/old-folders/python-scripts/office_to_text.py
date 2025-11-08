#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import sys
import json
import os

# Принудительно устанавливаем UTF-8 кодировку для stdout
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

def extract_text_from_excel(file_path):
    """Извлекает текст из Excel файлов (.xlsx, .xls)"""
    try:
        import pandas as pd
        import warnings
        
        # Подавляем warning сообщения от xlrd
        with warnings.catch_warnings():
            warnings.simplefilter("ignore")
            
            # Читаем все листы
            if file_path.lower().endswith('.xlsx'):
                sheets = pd.read_excel(file_path, sheet_name=None, engine='openpyxl')
            else:  # .xls
                # Для .xls файлов подавляем все выводы xlrd и пытаемся разные кодировки
                import os
                from contextlib import redirect_stdout, redirect_stderr
                with open(os.devnull, 'w') as devnull:
                    with redirect_stdout(devnull), redirect_stderr(devnull):
                        try:
                            # Пытаемся с Windows-1251 (русская кодировка)
                            import xlrd
                            book = xlrd.open_workbook(file_path, encoding_override='cp1251')
                            sheets = {}
                            for sheet_name in book.sheet_names():
                                sheet = book.sheet_by_name(sheet_name)
                                data = []
                                for row_idx in range(sheet.nrows):
                                    row_data = []
                                    for col_idx in range(sheet.ncols):
                                        cell_value = sheet.cell_value(row_idx, col_idx)
                                        row_data.append(cell_value)
                                    data.append(row_data)
                                df = pd.DataFrame(data)
                                sheets[sheet_name] = df
                        except:
                            # Если не сработало, используем стандартный способ
                            sheets = pd.read_excel(file_path, sheet_name=None, engine='xlrd')
        
        all_text = []
        
        for sheet_name, df in sheets.items():
            all_text.append(f"=== ЛИСТ: {sheet_name} ===")
            
            # Конвертируем DataFrame в текст построчно
            df_clean = df.fillna('')
            
            for _, row in df_clean.iterrows():
                row_text = []
                for value in row:
                    str_value = str(value).strip()
                    if str_value and str_value != 'nan' and str_value != 'None':
                        row_text.append(str_value)
                
                if row_text:
                    all_text.append(' '.join(row_text))
        
        return '\n'.join(all_text)
        
    except ImportError:
        return "Ошибка: Не установлена библиотека pandas или openpyxl/xlrd"
    except Exception as e:
        return f"Ошибка чтения Excel файла: {str(e)}"

def extract_text_from_word(file_path):
    """Извлекает текст из Word файлов (.docx, .doc)"""
    try:
        if file_path.lower().endswith('.docx'):
            # Используем python-docx для .docx
            from docx import Document
            
            doc = Document(file_path)
            text_parts = []
            
            # Извлекаем текст из параграфов
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Извлекаем текст из таблиц
            for table in doc.tables:
                text_parts.append("=== ТАБЛИЦА ===")
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_parts.append(' '.join(row_text))
            
            return '\n'.join(text_parts)
            
        else:  # .doc
            # Пытаемся использовать python-docx для старых форматов
            try:
                import docx2txt
                text = docx2txt.process(file_path)
                return text if text else "Не удалось извлечь текст из .doc файла"
            except ImportError:
                return "Ошибка: Для .doc файлов нужна библиотека docx2txt"
            
    except ImportError:
        return "Ошибка: Не установлена библиотека python-docx"
    except Exception as e:
        return f"Ошибка чтения Word файла: {str(e)}"

def main():
    try:
        if len(sys.argv) != 2:
            result = {
                "error": "Использование: python office_to_text.py <путь_к_файлу>",
                "text": "",
                "text_length": 0
            }
            print(json.dumps(result, ensure_ascii=False))
            sys.exit(1)
        
        file_path = sys.argv[1]
        
        if not os.path.exists(file_path):
            result = {
                "error": f"Файл не найден: {file_path}",
                "text": "",
                "text_length": 0
            }
            print(json.dumps(result, ensure_ascii=False))
            sys.exit(1)
        
        # Определяем тип файла
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension in ['.xlsx', '.xls']:
            text = extract_text_from_excel(file_path)
        elif file_extension in ['.docx', '.doc']:
            text = extract_text_from_word(file_path)
        else:
            result = {
                "error": f"Неподдерживаемый тип файла: {file_extension}. Поддерживаются: .xlsx, .xls, .docx, .doc",
                "text": "",
                "text_length": 0
            }
            print(json.dumps(result, ensure_ascii=False))
            sys.exit(1)
        
        # Проверяем, нет ли ошибки в тексте
        if text.startswith("Ошибка"):
            result = {
                "error": text,
                "text": "",
                "text_length": 0
            }
        else:
            result = {
                "text": text,
                "text_length": len(text),
                "file_path": file_path,
                "file_type": file_extension
            }
        
        # Выводим результат как JSON
        print(json.dumps(result, ensure_ascii=False))
        
    except Exception as e:
        # В случае любой ошибки возвращаем JSON с ошибкой
        error_result = {
            "error": f"Неожиданная ошибка: {str(e)}",
            "text": "",
            "text_length": 0
        }
        print(json.dumps(error_result, ensure_ascii=False))

if __name__ == "__main__":
    main()